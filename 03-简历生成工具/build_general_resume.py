from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/casperli/Documents/personal")
OUT_DIR = ROOT / "04-简历输出"
OUT_PATH = OUT_DIR / "Casper-通用版简历-可修改.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8BDC7", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_bottom_border(paragraph, color="B8BDC7", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def style_text(run, size=9.5, bold=False, color="222222"):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text, size=9.5, bold=False, color="222222"):
    run = paragraph.add_run(text)
    style_text(run, size=size, bold=bold, color=color)
    return run


def set_para(paragraph, before=0, after=2, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_section_title(document, title):
    paragraph = document.add_paragraph()
    set_para(paragraph, before=5, after=3, line=1)
    add_run(paragraph, title, size=11, bold=True, color="1F2937")
    set_bottom_border(paragraph)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_para(paragraph, before=0, after=1.2, line=1.03)
    add_run(paragraph, text, size=9.1)
    return paragraph


def add_experience_header(document, left, right):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(13.9)
    table.columns[1].width = Cm(3.1)
    set_cell_width(table.cell(0, 0), 13.9)
    set_cell_width(table.cell(0, 1), 3.1)
    for cell in table.row_cells(0):
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left_p = table.cell(0, 0).paragraphs[0]
    right_p = table.cell(0, 1).paragraphs[0]
    set_para(left_p, after=0, line=1)
    set_para(right_p, after=0, line=1)
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(left_p, left, size=9.6, bold=True)
    add_run(right_p, right, size=9.2, color="555555")
    return table


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def build_resume():
    OUT_DIR.mkdir(exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    bullet = styles["List Bullet"]
    bullet.font.name = "Aptos"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(9.1)
    bullet.paragraph_format.left_indent = Cm(0.45)
    bullet.paragraph_format.first_line_indent = Cm(-0.18)
    bullet.paragraph_format.space_after = Pt(1.2)
    bullet.paragraph_format.line_spacing = 1.03

    title = document.add_paragraph()
    set_para(title, after=1, line=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "李天译  Casper Li", size=19, bold=True, color="111827")

    subtitle = document.add_paragraph()
    set_para(subtitle, after=2, line=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        subtitle,
        "海外增长运营 | 数字营销 | 短剧整合营销",
        size=10,
        bold=True,
        color="374151",
    )

    contact = document.add_paragraph()
    set_para(contact, after=4, line=1)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        contact,
        "邮箱：lierhu0716@gmail.com  |  手机：待补充  |  微信：待补充  |  所在地：印尼",
        size=9.1,
        color="444444",
    )

    add_section_title(document, "个人概述")
    overview = document.add_paragraph()
    set_para(overview, after=2, line=1.06)
    add_run(
        overview,
        "具备海外短剧 App、内容增长与数字营销实战经验，当前聚焦东南亚市场。"
        "熟悉社媒买量、素材测试、内容矩阵、KOL/UGC 联动、社群反馈收集与转化漏斗分析，"
        "能够围绕曝光、获客、付费转化和运营效率推动增长动作落地。",
        size=9.3,
    )

    add_section_title(document, "核心能力")
    skill_table = document.add_table(rows=2, cols=3)
    skill_table.autofit = False
    skill_table.allow_autofit = False
    widths = [5.6, 5.6, 5.8]
    headers = ["增长与投放", "内容与传播", "数据与协作"]
    bodies = [
        "Meta 等渠道投放\n预算管理\n素材与转化测试",
        "短剧整合营销\nKOL / UGC / 社群\n本地化内容运营",
        "漏斗分析\n用户反馈提炼\nAI 工作流提效",
    ]
    for idx, width in enumerate(widths):
        set_cell_width(skill_table.cell(0, idx), width)
        set_cell_width(skill_table.cell(1, idx), width)
        for row in range(2):
            cell = skill_table.cell(row, idx)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        p = skill_table.cell(0, idx).paragraphs[0]
        set_para(p, after=0, line=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, size=9, bold=True, color="1F2937")
        set_cell_shading(skill_table.cell(0, idx), "EEF2F7")
    for idx, body in enumerate(bodies):
        p = skill_table.cell(1, idx).paragraphs[0]
        set_para(p, after=0, line=1.02)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, body, size=8.8, color="30343B")
    set_table_borders(skill_table)

    add_section_title(document, "工作经历")
    add_experience_header(
        document,
        "PT. Bintang Seni Media | 海外营销推广运营 | DopReel 东南亚短剧 App",
        "2025.07 - 至今",
    )
    add_bullet(
        document,
        "负责安卓端海外社媒买量投放与预算管理，覆盖核心渠道放量和日常投放优化；"
        "单端日消耗稳定 1,000-4,000 美元，产品整体日消耗峰值 10,000 美元。",
    )
    add_bullet(
        document,
        "结合社群反馈、投放数据与营销漏斗分析，推动 V2.1.0 付费墙展示逻辑优化，"
        "项目 ROAS 同比提升 45%。",
    )
    add_bullet(
        document,
        "参与短剧上线期整合营销，联动内容矩阵、KOL、社群与 UGC；"
        "《三年之约》上线期产出 10 条以上百万曝光内容，社群互动超过 1,000 条。",
    )

    add_experience_header(
        document,
        "网易有道 | 小红书获客运营",
        "2024.12 - 2025.05",
    )
    add_bullet(
        document,
        "管理投放账户并搭建聚光投放计划，优化课程获客留资路径；"
        "月精准留资线索 500+，日耗 1W+，LTV 1600，ROI 4。",
    )
    add_bullet(
        document,
        "拆解人群关注点并筛选高潜关键词，引入数字人素材生产 SOP；"
        "沉淀关键词 300+，曝光环比增长 60%，视频素材生产效率提升 400%。",
    )
    add_bullet(
        document,
        "针对高点击低留资素材进行对比测试和转化排查，推动获客效率提升 40%，"
        "转化率提升 19%。",
    )

    add_experience_header(
        document,
        "北京好未来教育集团 | TikTok 短视频运营",
        "2024.09 - 2024.12",
    )
    add_bullet(
        document,
        "负责海外 TikTok 双店铺选题、脚本、日更和复盘，产出爆款成交视频 10+；"
        "平均周 GMV 2,000 美元，环比增长 40%，ROI 稳定 12+。",
    )
    add_bullet(
        document,
        "完成短视频脚本、拍摄、包装和多平台分发，单条视频最高曝光 50W+；"
        "整合 AIGC 工具进入内容流程，制作效率提升 60%，拍摄成本减少 80%。",
    )
    add_bullet(
        document,
        "诊断独立站引流流失问题并设计短链跳转方案，引流路径缩短 70%，"
        "30 天用户转化率提升 25%。",
    )

    add_section_title(document, "代表项目")
    add_experience_header(
        document,
        "DopReel 付费墙转化优化 | 用户反馈 + 投放数据 + 营销漏斗分析",
        "商业化增长",
    )
    add_bullet(
        document,
        "从社群反馈与转化数据中定位付费展示问题，向产品优化提供依据并推动版本协作，"
        "支撑付费转化效率提升。",
    )

    add_section_title(document, "教育背景")
    add_experience_header(
        document,
        "山西大同大学 | 网络与新媒体 | 本科",
        "2021.09 - 2025.07",
    )

    add_section_title(document, "补充信息")
    supplement = document.add_paragraph()
    set_para(supplement, after=0, line=1.03)
    add_run(
        supplement,
        "语言：中文、英文  |  工具：广告投放平台、内容运营工具、AIGC 工具  |  "
        "作品集 / LinkedIn：待补充",
        size=9.1,
    )

    for table in document.tables:
        if table is not skill_table:
            remove_table_borders(table)

    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_resume()
